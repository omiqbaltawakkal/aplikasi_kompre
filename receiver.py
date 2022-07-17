# import the socket module

import re
import struct
import json
import socket
import tkinter.scrolledtext as scrolledtext
import base64
from threading import Thread
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import xlsxwriter
from tkinter import *
from tkinter import ttk
import pandas as pd


break_condition = True


window_width = 500
window_height = 200


class GUI():
    def __init__(self, root):
        self.running = 0  # not listening
        self.addr = None
        self.conn = None
        self.serverSocket = None

        self.server_ip = "10.233.79.249"
        self.server_port = 50392

        self.collection_data_value_only = list()
        self.header_list = list()
        
        self.excel_collection_value = list()

        self.screen_width = root.winfo_screenwidth()
        self.screen_height = root.winfo_screenheight()
        # find the center point
        self.center_x = int(0.2 * self.screen_width)
        self.center_y = int(0.2 * self.screen_height)

        # set the position of the window to the center of the screen
        #            width x height
        root.geometry(f'1000x200+{self.center_x}+{self.center_y}')

        self.frame = Frame(root)
        self.frame.pack(side=LEFT, anchor=N)

        self.ip_port = Label(self.frame, text="IP/PORT : ")
        self.ip_port.pack(side=LEFT, anchor=SW)

        self.ip_str_var = StringVar()
        self.ip_port_entry = Entry(self.frame, textvariable=self.ip_str_var)
        self.ip_port_entry.pack(side=LEFT, anchor=N)
        
        self.max_number = Label(self.frame, text="Max Size : ")
        self.max_number.pack(side=LEFT, anchor=SW)

        self.max_number_var = IntVar()
        self.max_number_entry = Entry(self.frame, textvariable=self.max_number_var)
        self.max_number_entry.pack(side=LEFT, anchor=N)
        
        self.startb = Button(self.frame, text="Start", command=self.startc)
        self.startb.pack(side=LEFT, anchor=N)

        self.generate = Button(self.frame, text="Generate", command=self.generateKKPAExcel)
        self.generate.pack(side=LEFT, anchor=N)

        # self.connectionl = Label(self.frame, text="Not Started")
        # self.connectionl.pack(side=LEFT, anchor=SW)
        self.textboxes = scrolledtext.ScrolledText(root, undo=True)
        self.textboxes.pack(expand=True, fill='both')
        self.addToTextbox("Not Started")
    
    def addToTextbox(self, message):
        self.textboxes.config(state=NORMAL)
        self.textboxes.insert(END, message +"\n")
        self.textboxes.config(state=DISABLED)
        
    def generateKKPAExcel(self):
        data = None
        if (self.excel_collection_value):
            data = self.excel_collection_value
        else: 
            f = open('readme_self.excel_collection_value.txt', 'r')
            data = f.read()
            data = json.loads(data)
            
        #asumsi end process = 1 uker -> data[0].get("kode_uker") == data[N].get("kode_uker")
        kode_uker = data[0].get("kode_uker")
        
        list_temp = list()
        for each in data :
            temp = dict()
            for key, value in each.items():
                strings = ''
                if type(value) == list:
                    if key.__contains__("appli"):
                        value = ', '.join(value)
                    else:
                        repl = r'\1'
                        if key.__contains__("disk"):
                            repl = r'\1 '
                        for item in value:
                            for keys, values in item.items():
                                values = re.sub('(\d+(\.\d+)?)', repl, values)
                                item.update({keys: values.strip()})
                            strings +='\n'.join(': '.join((key,val)) for (key,val) in item.items())
                            strings += "\n\n"
                        value = strings
                else:
                    if key.__contains__('processor'):
                        splits = value.strip().split("/")
                        number = re.findall(r'[A-Za-z]+|\d+(?:\.\d+)?', splits[0].split("@")[1])
                        strings += "Frekuensi: "+ " ".join(number)+ "\n"
                        strings += "Utilisasi: "+ splits[1] + "\n"
                        strings +="\n\n"
                        value = strings
                    elif key.__contains__("ram"):
                        splits = value.split("/")
                        strings += "Total: "+ splits[0] + "\n"
                        strings += "Utilisasi: "+ splits[1] +"\n"
                        strings +="\n\n"
                        value = strings
                    elif key.__contains__("address"):
                        splits = value.split("/")
                        value = splits[0]
                    elif key.__contains__("saver"):
                        value = value.get("Status")
                    elif key.__contains__("operating"):
                        value = '\n'.join(': '.join((key,val)) for (key,val) in value.items())
                    elif key == "ntp":
                        value = value.split("-")[0]
                    else:
                        pass
                temp.update({key.replace("_"," ").upper(): value})
            list_temp.append(temp)
        data = list_temp
        
        df = pd.DataFrame(data=data)
        # writer = pd.ExcelWriter('kkpa.xlsx', engine='xlsxwriter')
        writer = pd.ExcelWriter('KKPA.xlsx', engine='xlsxwriter')
        df.to_excel(writer, index=False, startrow=2, sheet_name='Sheet1')
        worksheet = writer.sheets['Sheet1']
        worksheet.write('A1', "Unit Kerja : " + kode_uker)
        writer.save()

    def generateKKPAFromRawData(self, client_raw_data):
        self.addToTextbox("generate KKPA start ...")
        dict_raw_data = client_raw_data
        document = Document()
        owner_name = dict_raw_data.get('nama')
        # Header
        heading_style = document.styles['Body Text']
        head = document.add_paragraph(style=heading_style).add_run(
           "KKPA Pengambilan Data PC " +owner_name)
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        head.font.size = Pt(20)
        head.font.bold = True
        image_width = Cm(15)
        
        exception_list = ["tanggal", "nama", "pn", "jabatan", "kode_uker"]
        # Template docx
        # informasi umum
        document.add_heading("Infomasi Umum", 1)
        for key, value in dict_raw_data.items():
            if str(key) in exception_list:
                doc = document.add_paragraph("")
                doc.add_run(
                    key.replace("_", " ").upper()).bold = True
                doc.add_run(" : {} ".format(value))
        p = doc.add_run()
        p.add_break()
        #spesifikasi 
        document.add_heading("Spesifikasi",1)
        for key, value in dict_raw_data.items():
            if type(value) == list:
                p = document.add_paragraph()
                p.add_run("{}".format(
                    key.replace("_", " ").upper())).bold = True
                if str(key).__contains__("disk"):
                    for item in value:
                        percent = int(int(item.get('FreeSpace'))/int(item.get('Size')))
                        document.add_paragraph(" Disk \"{}\" Total / Free Space : {} / {} / {}%".format(
                            item.get('Name'), item.get('Size'), item.get('FreeSpace'), percent), style='List Number')
                elif str(key).__contains__("antivirus"):
                    for item in value:
                        document.add_paragraph("{}, updated on {}".format(
                            item.get('Name'), item.get('Last Update')), style='List Number')
                elif str(key).__contains__("ip_addre"):
                    for item in value:
                        document.add_paragraph(
                            "{} ({}) ".format(item[0], item[1]), style='List Number')
                elif str(key).__contains__("appli"):
                    for item in value:
                        document.add_paragraph("{}".format(item), style='List Bullet')
                else:
                    for item in value:
                        document.add_paragraph("", style='List Number')
                        for key, value in item.items():
                            document.add_paragraph(" {} : {} ".format(
                                key, str(value)), style='List Bullet 2')
            else:
                if str(key) not in exception_list and not str(key).__contains__("saver") and not str(key).__contains__("image") and not str(key).__contains__("tambahan") and not str(key).__contains__("operating"):
                    doc = document.add_paragraph("")
                    doc.add_run(
                        key.replace("_", " ").upper()).bold = True
                    doc.add_run(" : {} ".format(value))
                elif str(key).__contains__("operating"):
                    doc = document.add_paragraph("")
                    doc.add_run(
                        key.replace("_", " ").upper()).bold = True
                    doc.add_run(" : {}, updated on {} ".format(value.get("Name"), value.get("Last Update Installation Date")))
                elif str(key).__contains__("saver"):
                    paragraph = ""
                    if value.get("Status") == "Active":
                        paragraph = "Desktop name {}, have screensaver activated with timeout {} seconds".format(
                                value.get('Name'), value.get('ScreenSaverTimeout'))
                    else:
                        paragraph = "No Active Screensaver"
                    doc = document.add_paragraph("")
                    doc.add_run(
                        key.upper()).bold = True
                    doc.add_run(": {}".format(paragraph))
        
        document.add_heading("Foto",1)
        for key, value in dict_raw_data.items():
            if str(key).__contains__("image"):
                image_name = owner_name + '_' + str(key)+'.jpg'
                myimage = base64.b64decode(value)
                f_image = open(image_name, 'wb')
                f_image.write(myimage)
                f_image.close()
                document.add_paragraph(
                    ("Foto " + " ".join(str(key).split("_")[1:])).title(), style="List Number")
                document.add_picture(image_name, width= image_width)
        document.add_heading("Informasi Tambahan",1)
        document.add_paragraph(dict_raw_data.get("informasi_tambahan"))
        #end template doc
        document.save("KKPA " + owner_name + ".docx")

    def socket_thread(self):
        self.addToTextbox("thread started..")
        self.excel_collection_value.clear()
        self.server_ip = str(self.ip_port_entry.get().split(":")[0])
        self.server_port = int(self.ip_port_entry.get().split(":")[1])
        self.max_number = int(self.max_number_entry.get())

        self.serverSocket = socket.socket(socket.AF_INET, socket.SOCK_STREAM)

        self.serverSocket.bind((self.server_ip, self.server_port))
        self.serverSocket.listen()
        self.addToTextbox("CONNECTED!")
        self.count = 0
        while(self.running == 1):
            (self.conn, self.addr) = self.serverSocket.accept()
            
            self.count = self.count + 1
            
            self.addToTextbox("Accepted {} connections so far, from {}".format(self.count, self.addr[0]))
        
            self.json_length = struct.unpack(">I", self.conn.recv(4))[0]
            self.json_data = b""
            while len(self.json_data) < self.json_length:
                self.chunk = self.conn.recv(min(40960, self.json_length - len(self.json_data)))
                if not self.chunk:
                    break
                self.json_data += self.chunk
            self.dic_data = json.loads(self.json_data.decode("utf-8"))
            # self.dic_data = json.loads(self.raw_data)
            self.generateKKPAFromRawData(self.dic_data)
            self.excel_collection_value.append(self.dic_data)
            with open('readme '+self.dic_data.get('nama')+'.txt', 'w') as f:
                f.write(str(self.dic_data))
                f.close()
            self.conn.sendall(str.encode("Received !!"))
            self.addToTextbox(f"Received from {self.addr[0]}!")
            self.textboxes.see(END)
            self.conn.close()
                # break
            if self.count == self.max_number:
                self.running = 0
                break
        if self.excel_collection_value:
            with open('readme_self.excel_collection_value.txt', 'w') as f:
                f.write(str(self.excel_collection_value))
                f.close()
        self.addToTextbox("Ended!")
        self.textboxes.see(END)
        self.generateKKPAExcel()
        self.serverSocket.close()
        root.destroy()

    def startc(self):
        if self.running == 0:
            self.addToTextbox("Starting thread")
            self.running = 1
            self.threads = Thread(target=self.socket_thread, daemon=True)
            self.threads.start()
            writer = pd.ExcelWriter('KKPA.xlsx', engine='xlsxwriter')
            writer.save()
        else:
            self.addToTextbox("thread already started.")

    def stopc(self):
        if self.running == 1:
            self.addToTextbox("stopping thread...")
            self.running = 0
            # self.threads.join()
        else:
            self.addToTextbox("thread not running...")


root = Tk()
root.title('Server-side PC Information Getter')
gui = GUI(root)
root.mainloop()
